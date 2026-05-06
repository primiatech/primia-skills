#!/usr/bin/env python3
"""
build_docx.py

Gera o relatório estratégico em .docx.

IMPORTANTE: Antes de rodar, leia /mnt/skills/public/docx/SKILL.md.

Toda a análise vem do JSON gerado pelo Claude. Este script só monta o documento.

Uso:
    python build_docx.py --output-dir output/ --analysis-json analysis.json --output report.docx
"""
import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_quote(doc, text, author=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'"{text}"')
    run.italic = True
    if author:
        run2 = p.add_run(f"  — {author}")
        run2.italic = True
        run2.font.size = Pt(9)


def add_image_safe(doc, image_path, width_cm=10):
    try:
        if Path(image_path).exists():
            doc.add_picture(str(image_path), width=Cm(width_cm))
            return True
    except Exception as e:
        print(f"  [erro imagem]: {e}", file=sys.stderr)
    return False


def build_report(channels, all_videos, top_videos_by_channel, analysis, output_dir, output_path):
    doc = Document()

    # === CAPA ===
    title = doc.add_heading("Análise Competitiva — YouTube", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    canais_str = ", ".join(c["title"] for c in channels)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(canais_str)
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().strftime("%d de %B de %Y"))

    doc.add_page_break()

    # === 1. SUMÁRIO EXECUTIVO ===
    add_heading(doc, "1. Sumário Executivo", level=1)
    add_paragraph(doc, analysis.get("sumario_executivo_intro",
        "Este relatório consolida a análise competitiva dos canais selecionados, "
        "extraindo padrões de conteúdo, voz do avatar e oportunidades de posicionamento."))

    decisoes = analysis.get("decisoes_recomendadas", [])
    if decisoes:
        add_paragraph(doc, "Decisões recomendadas (5 prioridades):", bold=True)
        for i, d in enumerate(decisoes[:5], 1):
            doc.add_paragraph(f"{i}. {d}", style="List Number")

    # === 2. METODOLOGIA ===
    add_heading(doc, "2. Metodologia", level=1)
    meta = analysis.get("metodologia", {})
    add_paragraph(doc, f"Modo de execução: {meta.get('modo', 'Padrão')}")
    add_paragraph(doc, f"Janela temporal: {meta.get('janela_dias', 180)} dias")
    add_paragraph(doc, f"Vídeos analisados (metadata): {meta.get('total_videos', len(all_videos))}")
    add_paragraph(doc, f"Vídeos com transcrição completa: {meta.get('total_transcricoes', '—')}")
    add_paragraph(doc, f"Comentários analisados: {meta.get('total_comentarios', '—')}")

    limitacoes = analysis.get("limitacoes", [])
    if limitacoes:
        add_paragraph(doc, "Limitações da coleta:", bold=True)
        for lim in limitacoes:
            doc.add_paragraph(lim, style="List Bullet")

    # === 3. PANORAMA DO NICHO ===
    add_heading(doc, "3. Panorama do Nicho", level=1)
    panorama = analysis.get("panorama", "")
    if panorama:
        add_paragraph(doc, panorama)

    # === 4. ANÁLISE POR CANAL ===
    add_heading(doc, "4. Análise por Canal", level=1)

    channel_analysis = analysis.get("channel_summaries", {})
    top_video_analysis = analysis.get("top_video_analysis", {})

    for i, ch in enumerate(channels, 1):
        slug = ch["slug"]
        s = channel_analysis.get(slug, {})

        add_heading(doc, f"4.{i} {ch['title']}", level=2)

        add_paragraph(doc, "Perfil estratégico", bold=True)
        add_paragraph(doc, s.get("perfil", "—"))

        if s.get("produto_ancora"):
            add_paragraph(doc, f"Produto âncora identificado: {s['produto_ancora']}", italic=True)

        if s.get("estagio_funil"):
            add_paragraph(doc, f"Estágio de funil dominante: {s['estagio_funil']}", italic=True)

        # Top 5 vídeos
        add_paragraph(doc, "Top vídeos", bold=True)
        tops = top_videos_by_channel.get(slug, [])[:5]
        for v in tops:
            ta = top_video_analysis.get(v["video_id"], {})
            sub = doc.add_paragraph()
            sub.add_run(f"• {v['title']}").bold = True
            doc.add_paragraph(
                f"  Views: {v['view_count']:,}  |  View ratio: {v['view_to_subscriber_ratio']}  |  "
                f"https://www.youtube.com/watch?v={v['video_id']}"
            )

            thumb = output_dir / slug / "thumbnails" / f"{v['video_id']}.jpg"
            add_image_safe(doc, thumb, width_cm=8)

            if ta.get("hook_decoded"):
                add_paragraph(doc, f"  Hook: {ta['hook_decoded']}")
            if ta.get("por_que_funcionou"):
                add_paragraph(doc, f"  Por que funcionou: {ta['por_que_funcionou']}")
            doc.add_paragraph()

    # === 5. MAPA DA VOZ DO AVATAR ===
    doc.add_page_break()
    add_heading(doc, "5. Mapa da Voz do Avatar", level=1)
    add_paragraph(doc, "Construído a partir da categorização dos comentários dos top vídeos. "
                       "As frases entre aspas são quotes literais — preserve a linguagem ao adaptar para copy.")

    voz = analysis.get("voz_avatar_estruturada", {})

    for sec_key, sec_title in [
        ("dores", "5.1 Dores recorrentes"),
        ("desejos", "5.2 Desejos"),
        ("objecoes", "5.3 Objeções"),
        ("perguntas", "5.4 Perguntas não-respondidas (= ideias de conteúdo)"),
        ("vocabulario", "5.5 Vocabulário do nicho"),
    ]:
        items = voz.get(sec_key, [])
        if not items:
            continue
        add_heading(doc, sec_title, level=2)
        for item in items:
            if isinstance(item, dict):
                titulo = item.get("titulo") or item.get("padrao") or ""
                quotes = item.get("quotes", [])
                freq = item.get("frequencia")
                p = doc.add_paragraph()
                run = p.add_run(f"• {titulo}")
                run.bold = True
                if freq:
                    p.add_run(f"  ({freq} menções)")
                for q in quotes[:3]:
                    add_quote(doc, q)
            else:
                doc.add_paragraph(f"• {item}")

    # === 6. ANÁLISE COMPARATIVA ===
    doc.add_page_break()
    add_heading(doc, "6. Análise Comparativa", level=1)

    matriz = analysis.get("comparativo_matrix", [])
    if matriz:
        add_heading(doc, "6.1 Matriz competitiva", level=2)
        headers = list(matriz[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in matriz:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row.get(h, ""))

    saturacao = analysis.get("saturacao", "")
    if saturacao:
        add_heading(doc, "6.2 Saturação (temas explorados por todos)", level=2)
        add_paragraph(doc, saturacao)

    gaps = analysis.get("gaps", [])
    if gaps:
        add_heading(doc, "6.3 Gaps (temas órfãos com demanda)", level=2)
        for g in gaps:
            if isinstance(g, dict):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(g.get("tema", ""))
                run.bold = True
                p.add_run(f" — {g.get('justificativa', '')}")
            else:
                doc.add_paragraph(g, style="List Bullet")

    # === 7. BANCO DE HOOKS ===
    add_heading(doc, "7. Banco de Hooks e Ângulos", level=1)
    add_paragraph(doc, "Estruturas decodificadas dos vídeos top — pra usar/adaptar nos próximos roteiros.")

    hooks = analysis.get("banco_de_hooks", [])
    for h in hooks:
        if isinstance(h, dict):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(h.get("tipo", ""))
            run.bold = True
            p.add_run(f": {h.get('exemplo', '')}")
            if h.get("origem"):
                p.add_run(f" (visto em: {h['origem']})").italic = True
        else:
            doc.add_paragraph(h, style="List Bullet")

    # === 8. RECOMENDAÇÕES ===
    doc.add_page_break()
    add_heading(doc, "8. Recomendações Estratégicas", level=1)

    recs = analysis.get("recomendacoes", [])
    for i, r in enumerate(recs, 1):
        add_heading(doc, f"Recomendação {i}: {r.get('o_que', '')}", level=2)
        add_paragraph(doc, "Por quê (evidência):", bold=True)
        add_paragraph(doc, r.get("por_que", ""))
        add_paragraph(doc, "Como executar:", bold=True)
        add_paragraph(doc, r.get("como", ""))

    # === 9. ANEXO ===
    doc.add_page_break()
    add_heading(doc, "9. Anexo", level=1)
    add_paragraph(doc, "IDs analisados:", bold=True)
    for ch in channels:
        slug = ch["slug"]
        tops = top_videos_by_channel.get(slug, [])
        doc.add_paragraph(f"{ch['title']} ({ch['channel_id']}):")
        for v in tops:
            doc.add_paragraph(f"  • {v['video_id']} — {v['title'][:80]}", style="List Bullet")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Relatório salvo em {output_path}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--analysis-json", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    base = Path(args.output_dir)

    channels_file = base / "channels.json"
    if not channels_file.exists():
        channels_file = base.parent / "channels.json"
    with open(channels_file, "r", encoding="utf-8") as f:
        channels = json.load(f)

    all_videos = []
    top_videos_by_channel = {}

    for ch in channels:
        slug = ch["slug"]
        ch_dir = base / slug
        if not ch_dir.exists():
            continue

        videos_file = ch_dir / "videos.json"
        if videos_file.exists():
            with open(videos_file, "r", encoding="utf-8") as f:
                videos = json.load(f)
            for v in videos:
                v["channel_title"] = ch["title"]
                all_videos.append(v)

        top_file = ch_dir / "top_videos.json"
        if top_file.exists():
            with open(top_file, "r", encoding="utf-8") as f:
                tops = json.load(f)
            for v in tops:
                v["channel_title"] = ch["title"]
            top_videos_by_channel[slug] = tops

    with open(args.analysis_json, "r", encoding="utf-8") as f:
        analysis = json.load(f)

    build_report(channels, all_videos, top_videos_by_channel, analysis, base, args.output)


if __name__ == "__main__":
    main()
