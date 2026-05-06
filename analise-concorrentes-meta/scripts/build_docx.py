#!/usr/bin/env python3
"""
build_docx.py — Constrói o relatório executivo em Word.

Uso:
    python build_docx.py \
        --enriched /home/claude/output/concorrente1/enriched_ads.json ... \
        --analysis /home/claude/output/analysis.json \
        --output /mnt/user-data/outputs/relatorio_concorrentes.docx

Expectativa do `analysis.json`:
{
    "title": "Análise Competitiva — Hotmart, Eduzz, Kiwify — Maio/2026",
    "executive_summary": ["bullet 1", "bullet 2", ...],
    "methodology": "texto sobre como foi coletado, período, limitações",
    "panorama": "texto descritivo + dados",
    "by_competitor": {
        "Hotmart": {
            "perfil": "...",
            "top_ads": [{"ad_id": "...", "screenshot": "/path/to.jpg",
                         "days_running": 90, "destaque": "..."}],
            "angulos": "...",
            "ofertas": "...",
            "funil": "..."
        }
    },
    "comparative_matrix": [...],   # mesmas colunas da xlsx
    "opportunities": ["..."],
    "recommendations": [{"titulo": "...", "evidencia": "...", "prioridade": "Alta"}],
    "appendix_ids": ["id1", "id2", ...]
}

Lembre-se: o Claude que usa esta skill deve ler /mnt/skills/public/docx/SKILL.md
ANTES de gerar o analysis.json e rodar este script, para seguir as boas práticas
oficiais de geração de Word.
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    """Adiciona tabela com cabeçalho estilizado."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Cabeçalho
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    # Linhas
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val) if val is not None else ""


def add_image_safe(doc, path, width_inches=4.5):
    """Adiciona imagem se existir; ignora silenciosamente se não."""
    try:
        if path and Path(path).exists():
            doc.add_picture(str(path), width=Inches(width_inches))
    except Exception as e:
        doc.add_paragraph(f"[imagem indisponível: {Path(path).name if path else '?'}]")


def build(analysis: dict, all_ads: dict, output: Path):
    doc = Document()

    # Configura estilo padrão
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Capa
    title = analysis.get("title", "Análise Competitiva — Meta Ad Library")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph()
    doc.add_paragraph()

    # 1. Resumo Executivo
    add_heading(doc, "1. Resumo Executivo", level=1)
    summary = analysis.get("executive_summary", [])
    if summary:
        add_bullet_list(doc, summary)
    else:
        add_paragraph(doc, "—")

    doc.add_page_break()

    # 2. Metodologia
    add_heading(doc, "2. Metodologia", level=1)
    add_paragraph(doc, analysis.get("methodology",
        "Coleta realizada via scraping da Biblioteca de Anúncios da Meta "
        "(interface pública). Dados extraídos: ID do anúncio, datas, plataformas, "
        "copy, CTA, mídia. Limitações: a Meta Ad Library não fornece métricas de "
        "performance (impressões, cliques, gasto) para anúncios comerciais — o "
        "principal proxy de performance utilizado é o tempo de veiculação "
        "(`days_running`)."))

    # 3. Panorama
    add_heading(doc, "3. Panorama Geral", level=1)
    add_paragraph(doc, analysis.get("panorama", "—"))

    # Tabela de panorama
    panorama_table = []
    for competitor, info in all_ads.items():
        ads = info["ads"]
        active = sum(1 for a in ads if a.get("is_active"))
        days = [a["days_running"] for a in ads if a.get("days_running") is not None]
        median = sorted(days)[len(days) // 2] if days else "—"
        panorama_table.append([competitor, len(ads), active, median])

    if panorama_table:
        add_table(doc,
                  ["Concorrente", "Total", "Ativos", "Mediana dias no ar"],
                  panorama_table)

    doc.add_page_break()

    # 4. Análise por Concorrente
    add_heading(doc, "4. Análise por Concorrente", level=1)

    by_comp = analysis.get("by_competitor", {})
    for i, (competitor, comp_data) in enumerate(by_comp.items(), 1):
        add_heading(doc, f"4.{i} {competitor}", level=2)

        if comp_data.get("perfil"):
            add_heading(doc, "Perfil de operação", level=3)
            add_paragraph(doc, comp_data["perfil"])

        if comp_data.get("top_ads"):
            add_heading(doc, "Top 5 anúncios mais longevos", level=3)
            for ad in comp_data["top_ads"][:5]:
                add_paragraph(doc,
                    f"ID {ad.get('ad_id')} — {ad.get('days_running', '?')} dias — "
                    f"{ad.get('destaque', '')}", bold=True)
                if ad.get("screenshot"):
                    add_image_safe(doc, ad["screenshot"], width_inches=3.5)
                doc.add_paragraph()

        if comp_data.get("angulos"):
            add_heading(doc, "Ângulos dominantes", level=3)
            add_paragraph(doc, comp_data["angulos"])

        if comp_data.get("ofertas"):
            add_heading(doc, "Ofertas recorrentes", level=3)
            add_paragraph(doc, comp_data["ofertas"])

        if comp_data.get("funil"):
            add_heading(doc, "Estrutura de funil", level=3)
            add_paragraph(doc, comp_data["funil"])

        doc.add_page_break()

    # 5. Análise Comparativa
    add_heading(doc, "5. Análise Comparativa", level=1)
    matrix = analysis.get("comparative_matrix", [])
    if matrix:
        headers = list(matrix[0].keys())
        rows = [[r.get(h, "") for h in headers] for r in matrix]
        add_table(doc, headers, rows)
    else:
        add_paragraph(doc, "—")

    # 6. Oportunidades
    add_heading(doc, "6. Oportunidades Identificadas", level=1)
    opps = analysis.get("opportunities", [])
    if opps:
        add_bullet_list(doc, opps)
    else:
        add_paragraph(doc, "—")

    # 7. Recomendações
    add_heading(doc, "7. Recomendações Estratégicas", level=1)
    recs = analysis.get("recommendations", [])
    for i, r in enumerate(recs, 1):
        add_paragraph(doc, f"{i}. {r.get('titulo', '')}", bold=True)
        if r.get("evidencia"):
            add_paragraph(doc, f"Evidência: {r['evidencia']}")
        if r.get("prioridade"):
            add_paragraph(doc, f"Prioridade: {r['prioridade']}")
        doc.add_paragraph()

    # 8. Anexo
    add_heading(doc, "8. Anexo — IDs Analisados", level=1)
    ids = analysis.get("appendix_ids", [])
    if ids:
        add_paragraph(doc, ", ".join(ids), size=9)
    else:
        all_ids = []
        for info in all_ads.values():
            all_ids.extend(a.get("ad_archive_id") for a in info["ads"] if a.get("ad_archive_id"))
        add_paragraph(doc, ", ".join(filter(None, all_ids)), size=9)

    doc.save(output)
    print(f"[✓] Relatório salvo em {output}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--enriched", nargs="+", required=True)
    parser.add_argument("--analysis", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    all_ads = {}
    for path in args.enriched:
        with open(path, encoding="utf-8") as f:
            d = json.load(f)
        all_ads[d["metadata"]["advertiser"]] = {"ads": d.get("ads", [])}

    with open(args.analysis, encoding="utf-8") as f:
        analysis = json.load(f)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    build(analysis, all_ads, output)


if __name__ == "__main__":
    main()
